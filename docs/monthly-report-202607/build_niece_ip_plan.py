# -*- coding: utf-8 -*-
"""侄女个人IP打造方案 docx"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

DOC_PATH = "/Users/xuhao/Projects/feixue-workshop/docs/monthly-report-202607/侄女个人IP打造方案.docx"

doc = Document()
for section in doc.sections:
    section.top_margin = Cm(2.6)
    section.bottom_margin = Cm(2.6)
    section.left_margin = Cm(2.6)
    section.right_margin = Cm(2.6)

DARK = RGBColor(0x33, 0x33, 0x33)
ACCENT = RGBColor(0xC0, 0x6B, 0x4D)
GREEN = RGBColor(0x2E, 0x7D, 0x32)


def set_font(run, name_cn, size, bold=False, color=None):
    run.font.name = name_cn
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    r = run._element.rPr.rFonts
    r.set(qn('w:eastAsia'), name_cn)


def add_para(text, font='仿宋', size=12, bold=False, align=None,
             indent=True, space_after=4, line=22, color=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    if align is not None:
        pf.alignment = align
    if indent:
        pf.first_line_indent = Pt(size * 2)
    pf.space_before = Pt(0)
    pf.space_after = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(line)
    run = p.add_run(text)
    set_font(run, font, size, bold, color)
    return p


def h1(text):
    add_para(text, font='黑体', size=14, bold=True, indent=False, space_after=6, line=26)

def h2(text):
    add_para(text, font='楷体', size=12, bold=True, indent=False, space_after=3, line=22)

def body(text):
    add_para(text, font='仿宋', size=12, indent=True, space_after=3, line=22)


# ===== 标题 =====
add_para('个人IP打造方案', font='黑体', size=20, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, space_after=4, line=30)
add_para('—— 给刚毕业、在昆明心霖公社（C&F School）工作的侄女 ——',
         font='楷体', size=11, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, space_after=12, line=22)

# ===== 一、总体定位 =====
h1('一、总体定位：一句话人设')
body('人设一句话："一个在昆明英语培训天花板机构实习的应届生，记录自己从学生到职场人的真实成长，顺带把英语学习的方法和工具分享给你。"')
body('三个关键词：真实（不装、不完美）、成长（正在发生的故事）、实用（背单词/刷题工具+方法）。')
body('核心逻辑：现在做IP不是马上变现，而是从入职第一天起就积累"作品集+人脉+话语权"。三年后无论她留在公社、跳槽、做自由老师还是留学申请，这个IP都是她最硬的简历。')

# ===== 二、平台矩阵与分工 =====
h1('二、平台矩阵：一主两翼')
h2('（一）个人网站（主阵地 · 资产沉淀）')
body('域名+网站是自己的资产，不依赖任何平台算法。放三样东西：')
body('1.成长博客：实习周记、工作思考、英语教学观察——按时间线积累，就是她的"数字人生档案"；')
body('2.小工具：背单词、刷题等在线小工具（详见第四部分）——工具是吸引流量的钩子，也是她技术能力的证明；')
body('3.关于我：人设页，讲清楚"我是谁、我在做什么、我能帮什么"。')
h2('（二）公众号（深度 · 私域沉淀）')
body('发长文：周记精选、教学方法思考、英语学习干货。公众号适合建立信任，是家长和学生群体最认的载体。频率：每周1篇，宁缺毋滥。')
h2('（三）小红书（流量 · 获客入口）')
body('发碎片：工作日常、办公室plog、英语学习技巧、夏令营现场。小红书是当下"大学生+教育"话题流量最大的平台，也是未来接家长咨询的主入口。频率：每周2-3条，轻量。')
body('三个平台内容同源、形式分拆：一篇周记 → 网站发全文、公众号发深度版、小红书拆3条碎片。一次生产，三处分发。')

# ===== 三、内容体系 =====
h1('三、内容体系：五大栏目')
h2('（一）实习/工作记录（人设基石）')
body('"第N天，我在心霖公社"系列：今天学了什么、带班遇到什么难题、被表扬/被批评了什么事。真实感是最大的护城河——现在网上全是精致人设，真实反而稀缺。')
h2('（二）英语教学观察（专业背书）')
body('以见习老师视角写：原来英语启蒙是这样的、一堂好课拆解、家长陪读的误区。这些内容直接服务公社的目标客群（家长），天然有受众。')
h2('（三）英语学习干货（引流内容）')
body('背单词方法、自然拼读、原版阅读书单、口语练习法。工具+方法组合拳，是全网最好传播的内容类型。')
h2('（四）小工具（差异化武器）')
body('详见第四部分。工具是"人无我有"的硬差异化，全网英语博主里会做工具的极少。')
h2('（五）个人成长反思（深度共鸣）')
body('从学生到职场人的心理转变、试用期的累与坚持、怎么跟同事领导相处。这类内容最容易引发同龄人共鸣，也最能体现思想深度。')

# ===== 四、小工具规划 =====
h1('四、小工具规划：背单词+刷题（差异化核心）')
h2('（一）为什么做工具')
body('工具解决三个问题：1.给用户一个"留下来"的理由（网站访问粘性）；2.证明她"英语+技术"的复合能力（简历加分项）；3.形成口碑传播（好用的工具会被主动分享）。')
h2('（二）首批三个工具（按难度排序）')
body('1.每日单词卡（最简单）：网页版单词卡片，每日10个，带发音、例句、记忆测试。可先用现成词库（中考/高考/四六级/雅思），后期支持自定义词单；')
body('2.英语晨读打卡：每天一段短文跟读+打卡日历，培养用户习惯，也培养她自己的内容习惯；')
body('3.小测刷题机（进阶）：选择题/完形填空题库，随机出题+错题本+正确率统计。题库可以从公开资料整理（注意版权，用自编或公开授权题目）。')
h2('（三）技术实现')
body('初期用静态网站即可（如 Next.js/Hexo/Hugo + 免费托管），她不需要会写复杂代码——可以先由我们这边搭建模板，她负责内容和维护。工具逻辑简单，背单词卡片就是"单词+发音+翻卡"，刷题机就是"题目+选项+判分"。')

# ===== 五、90天路线图 =====
h1('五、90天启动路线图（试用期阶段）')
h2('第1-2周：搭建地基')
body('注册域名（建议 hername.com 或 hername.cn）、搭建个人网站骨架、注册公众号+小红书账号、定好人设名和简介。此阶段不公开发内容，先想清楚再动。')
h2('第3-4周：内容试水')
body('公众号发第1篇（自我介绍+为什么做这件事），小红书发3-5条日常碎片，网站放1篇实习周记。观察反馈，调整方向。')
h2('第5-8周：形成节奏')
body('固定节奏：周记每周1篇（网站+公众号）、小红书每周2-3条、工具上线第1个（单词卡）。开始有意识收集工作素材（带班趣事、备课过程——注意脱敏）。')
h2('第9-12周：小闭环')
body('完成第2-3个工具，公众号粉丝过百/小红书有爆款苗头即可（不追求量）。把"试用期第100天"做成一个内容节点，回顾坚持的意义——这正好呼应她现在的累。')
body('节奏设计原则：每天投入不超过30分钟，绝不影响本职工作。试用期过关是第一优先级，IP是第二优先级。')

# ===== 六、合规与边界 =====
h1('六、合规与边界（重要）')
body('1.工作内容脱敏：写心霖公社的经历可以，但不能暴露学生姓名、家长信息、内部资料、薪资待遇。涉及具体教学场景，用"我们机构""我带的班"这种模糊表述；')
body('2.不蹭机构负面：不吐槽公司、不抱怨领导、不发"累到想辞职"这种内容——既保护自己，也保护她还在试用期的身份；')
body('3.肖像与隐私：不发学生正脸照片，发自己的照片注意别暴露机构内部敏感区域；')
body('4.版权：题库、文章引用注意来源，不自创题目或使用公开授权素材；')
body('5.定位明确：个人IP是"她自己的"，与公社无关，简介里不写"XX机构官方"，避免被解读为代表公司发声。')

# ===== 七、成本与分工 =====
h1('七、成本与分工')
body('成本：域名约60-100元/年，网站托管免费或极低（静态站），公众号/小红书免费。总启动成本控制在200元以内。')
body('分工建议：技术搭建（网站、工具）可由我们这边协助完成模板；内容创作（写作、拍照、选题）以她为主——IP的灵魂是她的真实生活，代笔可以搭框架，内容必须她自己来；运营策略（选题方向、数据复盘）定期给她指导。')
body('一句话收尾：现在种下的每一条记录、每一个工具，都是她未来最值钱的作品集。累的时候想想——她不是在打工，是在给三年后的自己写简历。')

doc.save(DOC_PATH)
print("saved:", DOC_PATH)
