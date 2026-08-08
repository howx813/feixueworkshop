# -*- coding: utf-8 -*-
"""生成 8月月度例会汇报 docx（3+5 格式，公文排版）"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

DOC_PATH = "/Users/xuhao/Projects/feixue-workshop/docs/monthly-report-202607/创新研究院8月月度例会汇报.docx"

doc = Document()
for section in doc.sections:
    section.top_margin = Cm(3.7)
    section.bottom_margin = Cm(3.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)


def set_font(run, name_cn, size, bold=False):
    run.font.name = name_cn
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)
    r = run._element.rPr.rFonts
    r.set(qn('w:eastAsia'), name_cn)


def add_para(text, font='仿宋', size=16, bold=False, align=None,
             indent=True, space_after=0, line=28):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    if align is not None:
        pf.alignment = align
    if indent:
        pf.first_line_indent = Pt(size * 2)
    pf.space_before = Pt(0)
    pf.space_after = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(line)
    run = p.add_run(text)
    set_font(run, font, size, bold)
    return p


# 标题
add_para('创新研究院2026年8月月度例会汇报', font='黑体', size=22, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, space_after=18, line=36)
add_para('（计划〔创新性〕工作3项＋任务〔事务性〕工作5项）', font='楷体', size=14,
         align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, space_after=24, line=28)

# 一、计划（创新性）工作3项
add_para('一、计划（创新性）工作3项', font='黑体', size=16, bold=True, indent=False)

add_para('（一）贵大数据平台招标项目投标', font='楷体', size=16, bold=True)
add_para('8月20日截止前，完成投标可行性评估与材料准备，启动标书编制。该项目预算约480万元，含8卡GPGPU计算节点与公共服务平台软件，是我院切入高校大数据公共服务平台赛道的关键一标，AI算法类业绩加分权重高，本月重点攻克。')

add_para('（二）商机雷达系统产出转化', font='楷体', size=16, bold=True)
add_para('推进商机雷达系统完成M4阶段建设，产出首期标讯周报，实现从"标讯自动采集"到"决策参考产出"的闭环，让AI工具直接服务经营决策。')

add_para('（三）省网安中心联合办学深化', font='楷体', size=16, bold=True)
add_para('组织人工智能训练师学员认定考试报名与考前准备，深化"招生代理+联合办学"合作模式，拓展培训生源与认证规模，形成可复制的培训收入模式。')

# 二、任务（事务性）工作5项
add_para('二、任务（事务性）工作5项', font='黑体', size=16, bold=True, indent=False)

add_para('（一）正高级工程师申报', font='楷体', size=16, bold=True)
add_para('推进副高取得年份、发明专利、标准等申报材料核实，8月28日前完成系统申报提交。')

add_para('（二）省重大专项实施', font='楷体', size=16, bold=True)
add_para('跟进已通过评审项目的8月批复，启动已批复2项省重大专项的实施工作。')

add_para('（三）TeleAgent推广落地', font='楷体', size=16, bold=True)
add_para('跟踪全省注册目标（预计282人）达成情况，推进SKILL技能凝练与应用场景打造，确保纳入科创月度画像考核。')

add_para('（四）自建系统优化迭代', font='楷体', size=16, bold=True)
add_para('末梢装维系统重点扩展至百富邦、盛通和等供应商使用；安管系统完成安全整改业务闭环功能并上线门户；全口径人资系统优化薪酬发放模块，根据创新院、工程公司试用情况改进软件。')

add_para('（五）网络信息安全保障', font='楷体', size=16, bold=True)
add_para('继续做好HW2026保障工作；完成XC云和数据库采购，按计划推进信创替代实施（预计10月完成）。')

doc.save(DOC_PATH)
print("saved:", DOC_PATH)
