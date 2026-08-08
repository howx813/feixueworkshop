# -*- coding: utf-8 -*-
"""通服机关食堂 8月3-7日 健康选餐卡片 docx"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOC_PATH = "/Users/xuhao/Projects/feixue-workshop/docs/monthly-report-202607/食堂8月3-7日健康选餐卡片.docx"

doc = Document()
for section in doc.sections:
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

GREEN = RGBColor(0x2E, 0x7D, 0x32)
AMBER = RGBColor(0xB8, 0x6E, 0x00)
RED = RGBColor(0xC6, 0x28, 0x28)
DARK = RGBColor(0x33, 0x33, 0x33)


def set_font(run, name_cn, size, bold=False, color=None):
    run.font.name = name_cn
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    r = run._element.rPr.rFonts
    r.set(qn('w:eastAsia'), name_cn)


def shade_cell(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)


def para(text, font='仿宋', size=11, bold=False, align=None, color=None, space_after=4):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    if align is not None:
        pf.alignment = align
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(0)
    pf.line_spacing = Pt(18)
    run = p.add_run(text)
    set_font(run, font, size, bold, color)
    return p


# 标题
para('通服机关食堂 · 健康选餐卡片', font='黑体', size=18, bold=True,
     align=WD_ALIGN_PARAGRAPH.CENTER, color=DARK, space_after=2)
para('2026年8月3日—8月7日 ｜ 按体检指标定制（血压偏高·甘油三酯↑·脂肪肝·胆囊息肉·肾结晶）',
     font='楷体', size=10, align=WD_ALIGN_PARAGRAPH.CENTER, color=AMBER, space_after=8)

# 三个全局原则
para('三条铁律', font='黑体', size=12, bold=True, color=DARK, space_after=2)
for t in ['① 少油：油炸、肥肉、浓油赤酱全回避（胆囊息肉+脂肪肝+甘油三酯）',
          '② 少盐：腌制品、腊味、卤味少吃（舒张压97偏高）',
          '③ 少嘌呤：浓肉汤少喝（肾结晶），例汤优先选青菜汤/绿豆汤/南瓜汤']:
    para(t, size=10.5, color=DARK, space_after=2)

para('', size=6, space_after=2)

# 表格
rows = [
    ('日期', '✅ 优先选', '⚠️ 少吃或换', '❌ 避开'),
    ('周一 8/3',
     '大虾、山药木耳甜肠(少量)、茄子豇豆、香菇菜心',
     '早餐脆哨/糟辣肉沫、麻辣豆腐丝、榨菜肉丝汤',
     '螺丝椒肉片(肥)、烫菜蘸水(辣油)'),
    ('周二 8/4',
     '芙蓉蛋、冬瓜、香菇、苹果',
     '折耳根拌鸡块、火腿烩冬瓜',
     '酸菜肉圆子汤(油大)；绿豆汤可喝'),
    ('周三 8/5',
     '卤龙骨(少量)、千页豆腐炒肉、尖椒小瓜丝、香梨',
     '蒜苗回锅肉(必肥)',
     '西红柿排骨汤→换素白萝卜汤'),
    ('周四 8/6',
     '萝卜烧牛腩、莴笋炒肉、西红柿烩土豆、西瓜',
     '干锅花菜(油重)、丝瓜肉片汤',
     '糟辣大白菜(咸辣)'),
    ('周五 8/7',
     '酸菜豆腐鱼、农家小炒肉(少量)、蒜苔肉沫、烧椒皮蛋、蒜泥空心菜',
     '泡椒魔芋、萝卜龙骨汤→换南瓜汤',
     '浓汤类一律换素汤'),
]

table = doc.add_table(rows=len(rows), cols=4)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

header_texts = rows[0]
for j, txt in enumerate(header_texts):
    cell = table.cell(0, j)
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(txt)
    set_font(run, '黑体', 10.5, True, RGBColor(0xFF, 0xFF, 0xFF))
    shade_cell(cell, '37474F')

colors = {1: GREEN, 2: AMBER, 3: RED}
for i in range(1, len(rows)):
    for j, txt in enumerate(rows[i]):
        cell = table.cell(i, j)
        cell.text = ''
        p = cell.paragraphs[0]
        if j == 0:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(txt)
        set_font(run, '仿宋', 9.5, j == 0, colors.get(j, DARK))
    if i % 2 == 0:
        for j in range(4):
            shade_cell(table.cell(i, j), 'F5F5F5')

para('', size=6, space_after=2)

# 三个提醒
para('三个提醒', font='黑体', size=12, bold=True, color=DARK, space_after=2)
for t in ['① 主食：保持不吃米饭；早餐避开糍粑、甜酒汤圆、芝麻大饼、小麻圆（糖油炸弹），选蒸玉米/山药/红薯/白粥/绿豆粥',
          '② 水果：西瓜、哈密瓜、火龙果偏甜，一天一份就够，别当水喝',
          '③ 多喝水：每天2000ml以上，冲淡肾结晶']:
    para(t, size=10.5, color=DARK, space_after=2)

para('', size=4, space_after=0)
para('—— 吃好，也是工作的一部分 ——', font='楷体', size=9,
     align=WD_ALIGN_PARAGRAPH.CENTER, color=AMBER, space_after=0)

doc.save(DOC_PATH)
print("saved:", DOC_PATH)
