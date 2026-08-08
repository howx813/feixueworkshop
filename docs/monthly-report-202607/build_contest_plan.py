# -*- coding: utf-8 -*-
"""生成：2026通信行业职工技能竞赛AI赛项策划方案 docx"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

DOC_PATH = "/Users/xuhao/Projects/feixue-workshop/docs/monthly-report-202607/2026通信行业职工技能竞赛AI赛项策划方案.docx"

doc = Document()
for section in doc.sections:
    section.top_margin = Cm(3.0)
    section.bottom_margin = Cm(3.0)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)


def set_font(run, name_cn, size, bold=False):
    run.font.name = name_cn
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)
    r = run._element.rPr.rFonts
    r.set(qn('w:eastAsia'), name_cn)


def add_para(text, font='仿宋', size=14, bold=False, align=None,
             indent=True, space_after=6, line=26):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    if align is not None:
        pf.alignment = align
    if indent:
        pf.first_line_indent = Pt(size * 2)
    pf.space_before = Pt(0)
    pf.space_after = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(line)
    run = p.add_run(text)
    set_font(run, font, size, bold)
    return p


def h1(text):
    add_para(text, font='黑体', size=16, bold=True, indent=False, space_after=8, line=30)

def h2(text):
    add_para(text, font='楷体', size=14, bold=True, indent=False, space_after=4, line=26)

def body(text):
    add_para(text, font='仿宋', size=14, indent=True, space_after=4, line=26)


# ===== 标题 =====
add_para('2026年贵州技能大赛·通信行业职工职业技能竞赛', font='黑体', size=20, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, space_after=4, line=32)
add_para('AI赛项内容设计与参赛夺标策划方案', font='黑体', size=18, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, space_after=8, line=30)
add_para('（内部策划稿 · 2026年8月）', font='楷体', size=12,
         align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, space_after=14, line=24)

# ===== 一、背景与机会 =====
h1('一、背景与机会判断')
h2('（一）赛事脉络')
body('2025年，省通信管理局、省人社厅、省总工会联合举办"2025年贵州技能大赛—贵州省通信行业互联网行业职工职业技能竞赛"，设信息通信营业员、信息通信网络线务员、数字化解决方案设计师3个赛项，理论笔试30%+实际操作70%。个人赛各赛项第一名可向省总工会申报"贵州省五一劳动奖章"，优胜奖以上选手可晋升相应职业三级（高级）职业技能等级。')
body('2026年，省人社厅已印发《关于组织开展2026年全省行业职业技能竞赛的通知》（黔人社函〔2026〕12号），明确纳入行业赛统一命名为"2026年贵州技能大赛—XX行业（系统）XX职业（工种）竞赛"。同时，省大数据局等五部门已启动"第二届数智技术职业技能竞赛暨第四届全国人工智能应用技术技能大赛贵州选拔赛"，AI进入省级技能竞赛赛道已是明确趋势。')
h2('（二）核心判断')
body('通信行业赛2026年大概率增设AI相关赛项或AI内容模块。当前正处于赛项设置与技术规则编制的前置窗口期——谁能参与赛项内容设计，谁就掌握了规则的主动权。省邮电学校（校长杨李红）是行业培训与技能鉴定体系的重要节点，是切入赛项设计的天然桥梁。')

# ===== 二、总体目标 =====
h1('二、总体目标')
body('总目标：参与AI赛项内容设计，使赛项方向与我院技术积累深度耦合，并组建参赛队冲击个人赛第一名，争取申报贵州省五一劳动奖章。')
body('三级目标：')
body('1.规则层：进入赛项技术委员会/命题专家组，把AI赛项设计成"我院有积累、对手难复制"的方向；')
body('2.成绩层：个人赛冲第一（可申报五一劳动奖章），团体赛力争一等奖，多名选手获优胜奖以上并晋升职业技能等级；')
body('3.生态层：借赛项沉淀培训产品（赛前集训、题库、认证），把"办赛+培训+参赛"变成我院年度业务闭环。')

# ===== 三、策略主线 =====
h1('三、策略主线：三线并进')
h2('（一）线一：参与规则设计（上游卡位）')
body('1.通过邮电学校、通管局技能鉴定中心渠道，争取进入大赛技术委员会或命题专家组，参与技术工作文件与命题。')
body('2.推荐的AI赛项设计方向（结合我院积累）：')
body('——AI应用解决方案设计：面向通信行业真实场景（智能运维、智慧客服、营销支撑），基于大模型/智能体完成方案设计与原型搭建，考核需求理解、场景拆解、工具链应用与方案表达；')
body('——AI训练数据工程：数据采集、清洗、标注、质检全流程，考核数据规范与质量意识，呼应我省数据标注产业定位；')
body('——AI+行业场景创新：开放命题，引导选手结合本省产业做AI落地创新。')
body('3.设计话术（名正言顺）："对标国家职业标准、突出AI+通信应用落地、体现贵州产业特色"，此口径既符合赛事"公平公正公开"生命线，又天然向我院优势方向倾斜。')
h2('（二）线二：备赛夺标（下游兑现）')
body('1.选手池：从AI培训已通过学员、人工智能训练师学员、数据标注团队、研发团队中选拔，组1支个人赛队伍（≤7人）并力争团体赛2队名额。')
body('2.集训方案：理论（职业标准、技术规范、题库）+实操（模拟赛题、真实场景演练）双线，8月下旬启动，9月选拔赛前完成两轮模拟。')
body('3.保障：场地（我院基地/邮电学校）、设备、师资、经费单列，指定专人统筹。')
h2('（三）线三：生态转化（赛后放大）')
body('1.赛前：争取承办或协办赛前培训、选拔赛组织，向全省通信企业输出培训服务，扩大我院行业影响力；')
body('2.赛后：获奖选手宣传报道，为"贵州工匠"评选、职称晋升积累素材；将赛项沉淀为年度培训产品，与人工智能训练师培训、网安中心合作形成联动。')

# ===== 四、关键动作与时间表 =====
h1('四、关键动作与时间表（8—10月）')
body('8月上旬：蓝善根牵头拜会杨李红校长，摸清2026年赛项设置动向与技能鉴定中心工作节奏，建立对接机制；')
body('8月中旬：提交《AI赛项内容设计建议书》（本方案第三部分为内核），经邮电学校/鉴定中心通道报送大赛组委会；同步启动参赛队组建与集训；')
body('8月下旬：跟进赛项设置反馈，动态调整集训重点；完成选手选拔确认与报名材料准备；')
body('9月：参加选拔赛，跟踪命题方向，针对性补强；')
body('10月：决赛冲击个人赛第一、团体赛一等奖，争取五一劳动奖章申报资格。')

# ===== 五、给邮电学校的合作点 =====
h1('五、给邮电学校的合作点（换规则参与权的"对价"）')
body('1.联合申报赛项承办/协办，提升学校在行业技能竞赛体系中的分量；')
body('2.共同开发AI培训课程与题库，学校有场地生源、我院有内容与师资，形成可持续培训合作；')
body('3.帮助学校建设AI实训环境、开展师资培训，助力其专业升级。')

# ===== 六、风险与合规边界（务必坚守） =====
h1('六、风险与合规边界（务必坚守）')
body('1.红线：大赛通知明确"公平、公正、公开是竞赛生命线"，全过程有监督、追责机制。参与规则设计≠内定成绩，严禁泄题、裁判舞弊、利益输送——五一劳动奖章需公示，一旦出事不仅荣誉归零，还会连累单位。')
body('2.正当路径：规则参与（赛项建议、技术文件起草、标准对接）+ 实力备赛（提前训练、模拟演练）+ 培训输出（帮行业整体水平提升，顺带建立知名度）。三者叠加已足以形成显著优势，无需触碰红线。')
body('3.廉政：与邮电学校、鉴定中心往来礼品礼金零容忍，合作事项留痕、走正式协议。')
body('4.舆情：赛前不对外宣传"参与命题"身份，避免"既当裁判又当运动员"的质疑；赛后统一宣传口径。')

# ===== 七、资源与分工 =====
h1('七、资源与分工建议')
body('对外联络：蓝善根（邮电学校、技能鉴定中心、组委会对接）；')
body('技术设计：创新研究院（赛项内容设计、技术文件起草、题库支撑）；')
body('选手培训：创新研究院+省网安中心合作团队（训练师课程复用）；')
body('统筹协调：院领导牵头，纳入8—10月重点工作跟踪。')

doc.save(DOC_PATH)
print("saved:", DOC_PATH)
