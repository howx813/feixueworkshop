# -*- coding: utf-8 -*-
"""生成 2026年7月工作月报 docx（公文格式：黑体标题/仿宋正文/首行缩进）"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

DOC_PATH = "/Users/xuhao/Projects/feixue-workshop/docs/monthly-report-202607/创新研究院2026年7月工作月报.docx"

doc = Document()

# 页边距（公文标准：上3.7 下3.5 左2.8 右2.6）
for section in doc.sections:
    section.top_margin = Cm(3.7)
    section.bottom_margin = Cm(3.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)


def set_font(run, name_cn, size, bold=False):
    run.font.name = name_cn
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)
    r = run._element.rPr.rFonts
    r.set(qn('w:eastAsia'), name_cn)


def add_para(text, font='仿宋', size=16, bold=False, align=None,
             indent=True, space_after=0, line=28):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    if align is not None:
        pf.alignment = align
    if indent:
        pf.first_line_indent = Pt(size * 2)
    pf.space_before = Pt(0)
    pf.space_after = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(line)
    run = p.add_run(text)
    set_font(run, font, size, bold)
    return p


# 标题（二号 方正小标宋 → 用黑体代替）
add_para('创新研究院2026年7月工作月报', font='黑体', size=22, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, space_after=24, line=36)

# 一、科技创新工作
add_para('一、科技创新工作', font='黑体', size=16, bold=True, indent=False)

add_para('（一）科创指标画像', font='楷体', size=16, bold=True)
add_para('6月科创月度画像得分74.3分，全国排名第6。其中研发合规30分基础分满分，AI告警动态清零，承接省重大专项获得加分。')

add_para('（二）对外申报', font='楷体', size=16, bold=True)
add_para('1.省大数据局申报3项（已提交系统）：', bold=True)
add_para('——人工智能大模型场景建设方向：《面向多场景的自主智能体平台关键技术研究与应用》，覆盖企业数转、交通、园区三个场景，预计申请专项资金300万元；')
add_para('——人工智能人才培养（3人）：年薪30万元及以上且在我省连续缴纳社会保险满一年，按每人5万元标准补助，申请资金15万元；')
add_para('——信息服务业项目—人工智能典型应用场景培育方向：《基于多模态信息融合的人体行为和生命体征预警与防控服务平台》完成现场答辩，申请专项资金300万元。')
add_para('2.省科技厅省市联动项目申请1项：成果转化（看护方向，多模态信息融合人体行为与生命体征预警防控服务平台）完成系统填报，申请专项资金100万元。')
add_para('3.首版次软件：数据融合与分析平台V1.0入选2026年度贵州省首版次软件产品，免审即享奖补20万元（待拨付）。')
add_para('4.省重大专项：2026年新申报的2项已正式批复，1项通过评审（预计8月批复）。')
add_para('5.人才考核：百层次、千层次人才中期考核专家评审通过，待批复后申请百层次终止培育。')

add_para('（三）人才培训认证', font='楷体', size=16, bold=True)
add_para('1.AI培训：累计通过6人，第三期上报7人（全年指标20人），已组织第4期培训；')
add_para('2.XC培训：第一期通过8人，1人未通过，6人缺考（全年指标15人）。')

add_para('（四）对外交流', font='楷体', size=16, bold=True)
add_para('1.重庆邮电大学网络空间安全与信息法学院现场调研交流；')
add_para('2.南明区科技局现场项目讲解及政策解读。')

# 二、数字化转型工作
add_para('二、数字化转型工作', font='黑体', size=16, bold=True, indent=False)

add_para('（一）AI+专班建设', font='楷体', size=16, bold=True)
add_para('完成AI+专班成立发文通知，明确专班组织架构与工作职责，统筹推进全院AI+应用落地。')

add_para('（二）TeleAgent推广', font='楷体', size=16, bold=True)
add_para('完成TeleAgent推广指标分解，全省预计注册282人；推广方案通过党委会审议，已组织全省开展安装注册工作。')

add_para('（三）内控信息化支撑', font='楷体', size=16, bold=True)
add_para('支撑贵州公司内控信息化数据整理工作，封存组织推进初见成效：冻结组织压降8.5%，项目压降71.6%。')

add_para('（四）自建系统优化', font='楷体', size=16, bold=True)
add_para('1.末梢装维系统：持续优化迭代，一线使用率持续提升；')
add_para('2.安管小程序：完成现场检查流程优化，上线新版小程序；')
add_para('3.全口径人资系统：完成人员数据清理打标。')

# 三、网络信息安全工作
add_para('三、网络信息安全工作', font='黑体', size=16, bold=True, indent=False)
add_para('1.HW2026：完成备战准备工作，随时准备临战；')
add_para('2.XC替代：启动XC替代工作，预计10月完成。')

# 四、下月工作计划
add_para('四、下月工作计划', font='黑体', size=16, bold=True, indent=False)

add_para('（一）科技创新', font='楷体', size=16, bold=True)
add_para('1.贵大数据平台招标：完成投标可行性评估与材料准备（8月20日截止）；')
add_para('2.正高级工程师申报：推进材料核实，8月28日前完成申报；')
add_para('3.省网安中心合作：组织人工智能训练师学员认定考试报名与考前准备；')
add_para('4.省重大专项：跟进评审项目批复，启动已批复项目实施。')

add_para('（二）数字化转型', font='楷体', size=16, bold=True)
add_para('1.持续优化末梢装维系统，重点将系统扩展至百富邦、盛通和等供应商使用；')
add_para('2.全口径人资系统：重点优化薪酬发放模块，根据创新院、工程公司试用情况改进软件；')
add_para('3.安管系统：完成安全整改业务闭环功能并上线门户。')

add_para('（三）网络信息安全', font='楷体', size=16, bold=True)
add_para('1.继续做好HW2026保障工作；')
add_para('2.完成XC云和数据库采购。')

doc.save(DOC_PATH)
print("saved:", DOC_PATH)
